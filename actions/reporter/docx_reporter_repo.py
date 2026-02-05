# Copyright (c) 2025 Linx Software, Inc.
#
# xlin-sbom-analysis tool is licensed under Mulan PSL v2.

# You can use this software according to the terms and conditions of the Mulan PSL v2.
# You may obtain a copy of Mulan PSL v2 at:
#
# http://license.coscl.org.cn/MulanPSL2
#
# THIS SOFTWARE IS PROVIDED ON AN "AS IS" BASIS, WITHOUT WARRANTIES OF ANY KIND,
# EITHER EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO NON-INFRINGEMENT,
# MERCHANTABILITY OR FIT FOR A PARTICULAR PURPOSE.
# See the Mulan PSL v2 for more details.


import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from actions.reporter.reporter_toolkit import (
    analyze_licenses,
    conclude_repo_report,
    categorize_packages,
    replace_placeholders
)
from actions.license_helper import LICENSE_CATEGORY_DETAILS
from actions import ASSIST_DIR

def _generate_license_section_docx(doc, license_summary, categorized_dict):
    """
    生成许可证分析部分的DOCX内容

    Args:
        doc (Document): python-docx文档对象，用于添加内容
        license_summary (dict): 许可证摘要信息，包含许可证类型及其出现次数
        categorized_dict (dict): 按许可证类别分类的软件包字典，键为许可证类别，值为Package对象列表

    Returns:
        None: 该函数直接修改传入的doc对象，不返回任何值
    """

    if not license_summary:
        doc.add_paragraph("未发现许可证信息。")
        return

    # 创建从scancode_category到详细信息的映射
    category_to_detail = {item["scancode_category"]
        : item for item in LICENSE_CATEGORY_DETAILS}

    analysis = analyze_licenses(license_summary)
    doc.add_paragraph(f"该update源中的软件包包含：")

    # 生成类别统计
    for category, count in analysis["category_counts"].items():
        if category in category_to_detail:
            desc = category_to_detail[category]["description"]
            doc.add_paragraph(f"{count}种{desc}", style='圆点')

    doc.add_paragraph()

    # 生成许可证表格
    for category, pkgs in categorized_dict.items():
        if not pkgs or category not in category_to_detail:
            continue

        doc.add_paragraph(f"其中具备 {category} 许可证的软件包如下：")

        # 添加suggestion内容
        suggestion = category_to_detail[category]["suggestion"]
        doc.add_paragraph(f"{suggestion}")

        # 创建表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # 添加表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '软件包名'
        hdr_cells[1].text = '版本'
        hdr_cells[2].text = '许可证'

        # 应用表头样式
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.style = '表格标题'

        # 填充表格数据
        for pkg in pkgs:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{pkg.name}"
            row_cells[1].text = f"{pkg.version}.{pkg.release}"
            row_cells[2].text = pkg.license

            # 应用表格内容样式
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = '表格内容'

        doc.add_paragraph()


def _generate_vulnerability_table_docx(doc, packages):
    """
    生成漏洞信息表格并添加到DOCX文档中

    Args:
        doc (Document): python-docx文档对象，用于添加表格内容
        packages (list): 软件包列表，每个元素为包含漏洞信息的Package对象

    Returns:
        None: 该函数直接修改传入的doc对象，不返回任何值
    """

    # 创建表格
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # 添加表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '软件包名'
    hdr_cells[1].text = '版本'
    hdr_cells[2].text = '漏洞ID'
    hdr_cells[3].text = '漏洞评分标准'
    hdr_cells[4].text = '漏洞评分'
    hdr_cells[5].text = '修复版本'

    # 应用表头样式
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.style = '表格标题'

    # 填充表格数据
    for package in packages:
        for vuln in package.vulnerabilities:
            row_cells = table.add_row().cells
            row_cells[0].text = package.name
            row_cells[1].text = f"{package.version}.{package.release}"
            row_cells[2].text = vuln['id']
            row_cells[3].text = vuln.get('severity_type', 'N/A')
            row_cells[4].text = vuln.get('severity_level', 'N/A')
            row_cells[5].text = vuln.get('fixed', 'N/A')

            # 应用表格内容样式
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = '表格内容'

    doc.add_paragraph()


def _generate_scan_results_table(doc, packages, config):
    """
    生成扫描结果汇总表格并添加到DOCX文档中

    Args:
        doc (Document): python-docx文档对象，用于添加表格内容
        packages (list): 软件包列表，每个元素为包含漏洞信息的Package对象
        config (object): 配置对象，用于获取配置信息

    Returns:
        None: 该函数直接修改传入的doc对象，不返回任何值
    """

    vuln_debug = config.get("general", {}).get(
        'debug_mode', {}).get('vulnerability', {})

    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # 添加表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '软件包'
    hdr_cells[1].text = '版本'
    hdr_cells[2].text = '漏洞ID'
    hdr_cells[3].text = '许可证'

    # 应用表头样式
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.style = '表格标题'

    # 填充表格数据
    for package in packages:
        row_cells = table.add_row().cells
        row_cells[0].text = package.name
        row_cells[1].text = f"{package.version}.{package.release}"
        row_cells[2].text = (
            ''
            if vuln_debug.get('enabled')
            else ', '.join(vuln['id'] for vuln in package.vulnerabilities)
        )
        row_cells[3].text = package.license

        # 应用表格内容样式
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.style = '表格内容'

    doc.add_paragraph()

def generate_docx_report(repo_name, start_date, end_date, packages,
                         license_summary, licenses_pie_chart, config):
    """
    生成DOCX格式的扫描报告

    Args:
        repo_name (str): 仓库名称
        start_date (str): 扫描开始日期
        end_date (str): 扫描结束日期
        packages (list): 软件包列表，包含Package对象
        license_summary (dict): 许可证摘要信息，包含许可证类型及其出现次数
        licenses_pie_chart (str): 许可证分布图的文件路径
        config (object): 配置对象，包含报告生成相关配置信息

    Returns:
        tuple: 包含三个元素的元组
            - doc (Document): 生成的python-docx文档对象
            - summary (list): 报告摘要信息列表
            - result (str): 评估结果，可能为"同意引入"或"拒绝引入"
    """

    vuln_debug = config.get("general", {}).get(
        'debug_mode', {}).get('vulnerability', {})

    # 加载模板文档
    template_path = os.path.join(ASSIST_DIR, "repo_report_base.docx")
    doc = Document(template_path)

    # 定位到模板中的"2详细测试记录"部分
    found_section = False
    for i, paragraph in enumerate(doc.paragraphs):
        if "详细测试记录" in paragraph.text:
            found_section = True
            # 清除该标题后的所有内容
            while len(doc.paragraphs) > i + 1:
                p = doc.paragraphs[i + 1]
                p._element.getparent().remove(p._element)
            break

    if not found_section:
        # 如果找不到指定部分，在文档末尾添加
        doc.add_heading("详细测试记录", level=1)

    # 2.1 扫描结果
    doc.add_heading("扫描结果", level=2)
    _generate_scan_results_table(doc, packages, config)

    # 2.2 漏洞分析
    doc.add_heading("漏洞分析", level=2)

    # 判断是否所有软件包都未检测到漏洞
    if vuln_debug.get('enabled') or all(not package.vulnerabilities for package in packages):
        doc.add_paragraph(
            f"经扫描确认，{repo_name} 中的软件包未检测到已知安全漏洞。"
            "当前扫描结果基于最新漏洞数据库，表明该update源在安全性方面处于良好状态。"
        )
    else:
        doc.add_paragraph(
            f"经扫描确认，{repo_name} 中的软件包存在以下安全漏洞："
        )
        # 生成漏洞表格
        _generate_vulnerability_table_docx(doc, packages)

        doc.add_paragraph(
            f"完整漏洞特征数据详见分析记录目录。"
        )

    # 2.3 许可证合规性分析
    doc.add_heading("许可证合规性分析", level=2)

    if license_summary:
        # 插入许可证分布图
        if licenses_pie_chart and os.path.exists(licenses_pie_chart):
            doc.add_picture(licenses_pie_chart, width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加图注
        caption = doc.add_paragraph(f"图1：{repo_name} 许可证分布图")
        caption.style = "插图"

        # 生成许可证分析内容
        _generate_license_section_docx(
            doc, license_summary, categorize_packages(packages))
    else:
        doc.add_paragraph("目标未检测到许可证信息。")

    # 第五章: 总结
    doc.add_heading("总结", level=2)
    summary, result, suggestion = conclude_repo_report(
        packages,
        license_summary,
        config
    )
    doc.add_paragraph(f"综合评估 {repo_name} 的安全状况如下：")
    for summary_text in summary:
        doc.add_paragraph(summary_text, style='圆点')

    # 替换文档中的占位符（包括页眉）
    brief_summary = ''.join(f"{i+1}. {summary_text}\n" for i,
                            summary_text in enumerate(summary))

    replace_placeholders(doc, repo_name, start_date, end_date,
                         brief_summary, result, suggestion, config)

    return doc, summary, result


